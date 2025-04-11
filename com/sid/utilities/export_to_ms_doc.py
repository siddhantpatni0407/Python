import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def create_title_page(doc):
    logging.info("Creating title page...")
    title = doc.add_heading('Java Interview Questions', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('For Freshers to Experienced (0–10+ Years)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)

    doc.add_paragraph('\nAuthor: Your Name')
    doc.add_paragraph('Last Updated: April 2025')
    doc.add_paragraph('Version: v1.0')
    doc.add_page_break()

def add_table_of_contents(doc):
    logging.info("Adding Table of Contents...")
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Core Java',
        '    1.1 OOPs Concepts',
        '    1.2 String, StringBuilder, StringBuffer',
        '    1.3 Exception Handling',
        '    1.4 Collections Framework',
        '    1.5 Multithreading and Concurrency',
        '    1.6 Java 8/11/17+ Features',
        '    1.7 JVM Internals'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_page_break()

def add_question_section(doc):
    logging.info("Adding questions...")

    doc.add_heading('1. Core Java', level=1)

    # 1.1 OOPs Concepts
    doc.add_heading('1.1 OOPs Concepts', level=2)
    add_question(doc,
        "What is the difference between abstraction and encapsulation in Java?",
        "Abstraction is the process of hiding the implementation details and showing only the functionality to the user. "
        "Encapsulation is wrapping up data and methods into a single unit (class).",
        "- Abstraction is achieved using abstract classes and interfaces.\n"
        "- Encapsulation is achieved using access modifiers (private, public, protected)."
    )

    # 1.2 Collections Framework
    doc.add_heading('1.2 Collections Framework', level=2)
    add_question(doc,
        "What is the difference between HashMap and ConcurrentHashMap?",
        "HashMap is not thread-safe and allows one null key and multiple null values. "
        "ConcurrentHashMap is thread-safe and does not allow null keys or null values.",
        "- HashMap is suitable for single-threaded environments.\n"
        "- ConcurrentHashMap uses internal locking for thread safety."
    )

    # 1.3 Multithreading
    doc.add_heading('1.3 Multithreading and Concurrency', level=2)
    add_question(doc,
        "What is the difference between a process and a thread?",
        "A process is an independent executing program, while a thread is a subset of the process that can run concurrently.",
        "- Threads share memory within the same process.\n"
        "- Processes do not share memory by default."
    )

    # 1.4 Java 8 Features
    doc.add_heading('1.4 Java 8 Features', level=2)
    add_question(doc,
        "What are functional interfaces in Java 8?",
        "A functional interface is an interface that contains only one abstract method. They are used as the basis for lambda expressions.",
        "- Examples include Runnable, Callable, Comparator, and custom ones annotated with @FunctionalInterface."
    )

def add_question(doc, question, answer, insights):
    doc.add_paragraph(f'❓ Question:\n{question}', style='Normal')
    doc.add_paragraph(f'✅ Answer:\n{answer}', style='Normal')
    doc.add_paragraph(f'🔍 Additional Insights:\n{insights}', style='Normal')

def generate_docx():
    try:
        doc = Document()
        create_title_page(doc)
        add_table_of_contents(doc)
        add_question_section(doc)

        # Format the filename with date and timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Java-Interview-Questions-{timestamp}.docx"
        output_dir = "C:/export"
        output_path = os.path.join(output_dir, filename)

        os.makedirs(output_dir, exist_ok=True)
        doc.save(output_path)

        logging.info(f"Document saved successfully at: {output_path}")
        return output_path
    except Exception as e:
        logging.error("Failed to create document", exc_info=True)
        return None

# Run the document generator
if __name__ == "__main__":
    generate_docx()
